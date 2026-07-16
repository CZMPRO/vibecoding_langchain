"""多格式文档加载：把各种文件读成纯文本块。"""

from pathlib import Path
from typing import List

import pandas as pd
from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".csv", ".xlsx", ".xls"}


def load_file_to_documents(file_path: Path, filename: str) -> List[LCDocument]:
    """按后缀选择解析方式，返回 LangChain Document 列表。"""
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型: {suffix}")

    if suffix == ".pdf":
        return _load_pdf(file_path, filename)
    if suffix == ".docx":
        return _load_docx(file_path, filename)
    if suffix in {".txt", ".md", ".markdown"}:
        return _load_text(file_path, filename)
    if suffix == ".csv":
        return _load_csv(file_path, filename)
    if suffix in {".xlsx", ".xls"}:
        return _load_excel(file_path, filename)
    raise ValueError(f"未实现的解析: {suffix}")


def _load_pdf(path: Path, filename: str) -> List[LCDocument]:
    reader = PdfReader(str(path))
    docs: List[LCDocument] = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            docs.append(
                LCDocument(
                    page_content=text,
                    metadata={"filename": filename, "page": i + 1, "file_type": "pdf"},
                )
            )
    if not docs:
        raise ValueError("PDF 未提取到文字（可能是扫描件，暂不支持 OCR）")
    return docs


def _load_docx(path: Path, filename: str) -> List[LCDocument]:
    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(paragraphs)
    if not text:
        raise ValueError("Word 文档内容为空")
    return [LCDocument(page_content=text, metadata={"filename": filename, "file_type": "docx"})]


def _load_text(path: Path, filename: str) -> List[LCDocument]:
    raw = path.read_bytes()
    text = None
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = raw.decode("utf-8", errors="ignore")
    text = text.strip()
    if not text:
        raise ValueError("文本文件内容为空")
    ftype = "md" if path.suffix.lower() in {".md", ".markdown"} else "txt"
    return [LCDocument(page_content=text, metadata={"filename": filename, "file_type": ftype})]


def _load_csv(path: Path, filename: str) -> List[LCDocument]:
    try:
        df = pd.read_csv(path)
    except UnicodeDecodeError:
        df = pd.read_csv(path, encoding="gbk")
    return _dataframe_to_docs(df, filename, "csv")


def _load_excel(path: Path, filename: str) -> List[LCDocument]:
    df = pd.read_excel(path)
    return _dataframe_to_docs(df, filename, "xlsx")


def _dataframe_to_docs(df: pd.DataFrame, filename: str, file_type: str) -> List[LCDocument]:
    """表格按行转成文本，表头写进每行，方便检索商品参数。"""
    if df.empty:
        raise ValueError("表格没有数据")
    cols = [str(c) for c in df.columns.tolist()]
    docs: List[LCDocument] = []
    # 每 20 行合成一块，避免块过碎
    batch_size = 20
    rows = df.fillna("").astype(str).values.tolist()
    for start in range(0, len(rows), batch_size):
        chunk_rows = rows[start : start + batch_size]
        lines = []
        for row in chunk_rows:
            pairs = [f"{cols[i]}: {row[i]}" for i in range(len(cols))]
            lines.append("；".join(pairs))
        content = f"来源表格: {filename}\n" + "\n".join(lines)
        docs.append(
            LCDocument(
                page_content=content,
                metadata={
                    "filename": filename,
                    "file_type": file_type,
                    "row_start": start + 1,
                    "row_end": start + len(chunk_rows),
                },
            )
        )
    return docs
